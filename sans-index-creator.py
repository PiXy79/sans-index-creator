from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict
import os
import sys
from tabulate import tabulate

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

# Ensure Excel file exists
if not os.path.exists('index.xlsx'):
    print("Error: 'index.xlsx' not found. Place the file 'index.xlsx' in the current directory.")
    sys.exit(1)

# Load Excel file
wb = load_workbook('index.xlsx')
ws = wb.active


# Parse data: (label, page_ref)
index_entries = []

for row in ws.iter_rows(min_row=1, values_only=True):
    if row[0]:  # Assuming column A = label, column B = x.yyy format
        label = str(row[0]).strip()
        page_ref = str(row[1]).strip() if row[1] else ""
        index_entries.append((label, page_ref))

# Sort alphabetically
index_entries.sort(key=lambda x: x[0].lower())

# Group by first letter
grouped = defaultdict(list)
for label, page_ref in index_entries:
    first_letter = label[0].upper()
    if first_letter.isalpha():
        grouped[first_letter].append((label, page_ref))

# Print statistics
print("\n" + "="*50)
print("📊 INDEX GENERATION STATISTICS")
print("="*50)

# Total entries statistics
print(f"\n📈 Total Statistics:")
total_read = len(index_entries)
total_produced = sum(len(entries) for entries in grouped.values())
print(f"   • Entries read from Excel: {total_read}")
print(f"   • Entries in final index: {total_produced}")

# Page reference statistics
entries_with_page_ref = sum(1 for _, page_ref in index_entries if page_ref)
entries_without_page_ref = total_read - entries_with_page_ref
print(f"\n📄 Page Reference Statistics:")
print(f"   • With page references: {entries_with_page_ref}")
print(f"   • Without page references: {entries_without_page_ref}")

# Entries per letter statistics
print(f"\n📝 Entries per Letter:")
letter_stats = []
for letter in sorted(grouped.keys()):
    count = len(grouped[letter])
    letter_stats.append([letter, count])

table_data = sorted(letter_stats, key=lambda x: x[0])
print(tabulate(table_data, headers=["Letter", "Count"], tablefmt="simple"))

# Min/Max per letter
counts = [count for _, count in letter_stats]
min_count = min(counts) if counts else 0
max_count = max(counts) if counts else 0
min_letter = next(letter for letter, count in letter_stats if count == min_count)
max_letter = next(letter for letter, count in letter_stats if count == max_count)

# Entry length statistics
entry_lengths = [len(label) for label, _ in index_entries]
min_length = min(entry_lengths) if entry_lengths else 0
max_length = max(entry_lengths) if entry_lengths else 0
min_entry = next((label for label, _ in index_entries if len(label) == min_length), "")
max_entry = next((label for label, _ in index_entries if len(label) == max_length), "")

# Summary stats
total_letters = len(grouped)
avg_per_letter = total_produced / total_letters if total_letters > 0 else 0
print(f"\n📊 Summary:")
print(f"   • Total letters: {total_letters}")
print(f"   • Average entries per letter: {avg_per_letter:.2f}")
print(f"   • Most entries: {max_letter} ({max_count})")
print(f"   • Least entries: {min_letter} ({min_count})")
print(f"   • Shortest entry: '{min_entry}' ({min_length} chars)")
print(f"   • Longest entry: '{max_entry}' ({max_length} chars)")
print("="*50 + "\n")

# Create Word document
doc = Document()

# Set default font size to smaller
doc.styles['Normal'].font.size = Pt(9)
doc.styles['Normal'].font.name = 'Arial'

# Set heading font to Arial
doc.styles['Heading 1'].font.name = 'Arial'
doc.styles['Heading 1'].font.size = Pt(20)

# Set document to one column
section = doc.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '1')

# Set smaller margins
section.left_margin = Inches(1.0)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

# Add page numbers to footer
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = ""
footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_paragraph.style.font.size = Pt(6)
run = footer_paragraph.add_run()
add_page_number(run)

if not index_entries:
    doc.add_paragraph("No index entries found in index.xlsx. Please add data to the Excel file.")
    doc.save('SANS_Index.docx')
    print("No data found. Generated document with message.")
    exit()

# Add main index with letter sections
for letter in sorted(grouped.keys()):
    heading = doc.add_paragraph(letter)
    heading.style = 'Heading 1'
    # Adjust heading vertical spacing (controls the visual "height" of the letter heading)
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(6)
    heading_format.space_after = Pt(4)
    heading_format.line_spacing = 1.0
    
    table = doc.add_table(rows=1, cols=2)
    # Make the table span more of the page and give more space to the first column
    table.autofit = False
    table.width = Inches(6.0)
    table.columns[0].width = Inches(4.0)
    table.columns[1].width = Inches(2.0)
        
    # Add entries for this letter
    for label, page_ref in grouped[letter]:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = page_ref
        # Set cell margins to prevent text overflow
        for cell in row_cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            margins = {'top': 40, 'left': 100, 'bottom': 40, 'right': 100}
            for margin_name, margin_value in margins.items():
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), str(margin_value))
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)
            # Remove paragraph spacing
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
    
    # Set smaller row height but allow expansion for content
    for row in table.rows:
        row.height = Pt(8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

# Save document
doc.save('SANS_Index.docx')
print("✅ Index generated: SANS_Index.docx")
